###############################################################################
# app.py
###############################################################################
import os
import re
import time
import random
import uuid
import traceback
from io import BytesIO

from flask import Flask, send_from_directory, request, jsonify, send_file
from docx import Document

from openai import AzureOpenAI

# For reading images, DOCX, PDF (via Form Recognizer)
from vision_api import analyze_image_from_bytes
from document_processing import extract_text_from_docx
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential

# Azure Cosmos + SendGrid
from azure.cosmos import CosmosClient, PartitionKey, exceptions
import sendgrid
from sendgrid.helpers.mail import Mail, Email, To, Content

# Azure Cognitive Search
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential

# Azure Blob (temp container)
from azure.storage.blob import BlobServiceClient, ContentSettings

###############################################################################
# 1. Flask and Cosmos Setup
###############################################################################
app = Flask(__name__, static_folder="src/public", static_url_path="")

COSMOS_ENDPOINT = os.environ.get("COSMOS_ENDPOINT", "")
COSMOS_KEY = os.environ.get("COSMOS_KEY", "")
COSMOS_DATABASE_ID = "GYMAIEngineDB"
COSMOS_CONTAINER_ID = "chats"

try:
    cosmos_client = CosmosClient(COSMOS_ENDPOINT, credential=COSMOS_KEY)
    database = cosmos_client.create_database_if_not_exists(id=COSMOS_DATABASE_ID)
    container = database.create_container_if_not_exists(
        id=COSMOS_CONTAINER_ID,
        partition_key=PartitionKey(path="/userKey"),
        offer_throughput=400
    )
except exceptions.CosmosHttpResponseError as e:
    app.logger.error("Cosmos DB setup error: %s", e)

###############################################################################
# 2. SendGrid + Form Recognizer + Blob
###############################################################################
SENDGRID_API_KEY = os.environ.get("SENDGRID_API_KEY", "")
if not SENDGRID_API_KEY:
    app.logger.warning("Missing SENDGRID_API_KEY (sendgrid usage might fail).")

FORM_RECOGNIZER_ENDPOINT = os.environ.get("FORM_RECOGNIZER_ENDPOINT", "")
FORM_RECOGNIZER_KEY = os.environ.get("FORM_RECOGNIZER_KEY", "")

AZURE_STORAGE_CONNECTION_STRING = os.environ.get("AZURE_STORAGE_CONNECTION_STRING", "")
AZURE_TEMP_CONTAINER = os.environ.get("AZURE_TEMP_CONTAINER", "temp-uploads")

blob_service_client = None
temp_container_client = None
if AZURE_STORAGE_CONNECTION_STRING:
    try:
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        temp_container_client = blob_service_client.get_container_client(AZURE_TEMP_CONTAINER)
        try:
            temp_container_client.create_container()
            app.logger.info(f"Created container '{AZURE_TEMP_CONTAINER}'")
        except Exception as ex:
            if "ContainerAlreadyExists" in str(ex):
                app.logger.info(f"Container '{AZURE_TEMP_CONTAINER}' already existed; ignoring.")
            else:
                app.logger.error("Error creating container: %s", ex)
                temp_container_client = None
    except Exception as ex:
        app.logger.error("Error initializing BlobServiceClient: %s", ex)
        temp_container_client = None
else:
    app.logger.warning("No AZURE_STORAGE_CONNECTION_STRING found. Temp file uploads will fail.")

###############################################################################
# 3. AzureOpenAI Setup
###############################################################################
client = AzureOpenAI(
    api_key=os.environ.get("AZURE_OPENAI_KEY"),
    azure_endpoint=os.environ.get("AZURE_OPENAI_ENDPOINT"),
    api_version="2023-05-15"
)
AZURE_DEPLOYMENT_NAME = "GYMAIEngine-gpt-4o"

###############################################################################
# 4. In-Memory Cache for Generated Reports
###############################################################################
report_cache = {}

###############################################################################
# 5. Helper Functions
###############################################################################
def extract_text_from_pdf(file_bytes):
    if not FORM_RECOGNIZER_ENDPOINT or not FORM_RECOGNIZER_KEY:
        raise ValueError("FORM_RECOGNIZER_ENDPOINT and FORM_RECOGNIZER_KEY must be set.")
    try:
        doc_client = DocumentAnalysisClient(
            endpoint=FORM_RECOGNIZER_ENDPOINT,
            credential=AzureKeyCredential(FORM_RECOGNIZER_KEY)
        )
        poller = doc_client.begin_analyze_document("prebuilt-document", file_bytes)
        result = poller.result()
        all_text = []
        for page in result.pages:
            for line in page.lines:
                all_text.append(line.content)
        return "\n".join(all_text)
    except Exception as ex:
        app.logger.error("PDF extraction failed:", exc_info=True)
        raise RuntimeError(f"Error extracting PDF text: {ex}")

def generate_detailed_report(base_content):
    detail_messages = [
        {
            "role": "system",
            "content": (
                "You are a helpful assistant that specializes in creating detailed, "
                "comprehensive reports."
            )
        },
        {
            "role": "user",
            "content": (
                "Here is a brief summary: " + base_content + "\n\n"
                "Please create a significantly more in-depth and expanded report."
            )
        }
    ]
    try:
        response = client.chat.completions.create(
            messages=detail_messages,
            model=AZURE_DEPLOYMENT_NAME
        )
        return response.choices[0].message.content
    except Exception as e:
        app.logger.error("Error calling AzureOpenAI for detailed report:", exc_info=True)
        return base_content + "\n\n(Additional detail could not be generated.)"

###############################################################################
# 6. Serve Frontend
###############################################################################
@app.route("/")
def serve_frontend():
    return send_from_directory("src/public", "index.html")

###############################################################################
# 7. Chat Endpoint (Now using create_item / replace_item instead of upsert)
###############################################################################
@app.route("/chat", methods=["POST"])
def chat_endpoint():
    """
    Creates a new chat doc if chatId is not provided (with a truly unique ID).
    If chatId is provided, we attempt to read that doc and then replace it
    upon changes. No 'upsert_item' is used here.
    """
    user_key = request.args.get("userKey", "default_user")
    chat_id = None
    user_input = ""
    uploaded_files = []

    # Grab request content
    if request.content_type and "multipart/form-data" in request.content_type:
        user_input = request.form.get("userMessage", "")
        chat_id = request.form.get("chatId")
        uploaded_files = request.files.getlist("file") or []
    else:
        data = request.get_json(force=True) or {}
        user_input = data.get("userMessage", "")
        user_key = data.get("userKey", user_key)
        chat_id = data.get("chatId")
        uploaded_files = []

    # 1) Guarantee uniqueness if no chatId
    is_new_doc = False
    if not chat_id:
        while True:
            temp_id = f"chat_{int(time.time()*1000)}_{random.randint(1000,9999)}_{user_key}"
            existing = list(container.query_items(
                query="SELECT * FROM c WHERE c.id=@id",
                parameters=[{"name": "@id", "value": temp_id}],
                enable_cross_partition_query=True
            ))
            if len(existing) == 0:
                chat_id = temp_id
                is_new_doc = True
                break

    # Prepare base messages for AzureOpenAI
    messages_for_openai = [
        {
            "role": "system",
            "content": (
                "You are a helpful assistant. Use Markdown. "
                "If user requests a downloadable report, MUST provide `download://report.docx`. "
                "If references, show them, else `References: None`."
            )
        },
        {"role": "user", "content": user_input}
    ]

    # 2) Either read existing doc or create from scratch
    chat_doc = None
    if not is_new_doc:
        # Means user supplied a chat_id, or we discovered one
        try:
            existing_doc = container.read_item(item=chat_id, partition_key=user_key)
            chat_doc = existing_doc
        except exceptions.CosmosResourceNotFoundError:
            # The user gave a chatId but it doesn't exist => We'll create it from scratch
            chat_doc = {
                "id": chat_id,
                "userKey": user_key,
                "messages": [],
                "files": []
            }
            is_new_doc = True
    else:
        # We have a brand-new ID => brand-new doc
        chat_doc = {
            "id": chat_id,
            "userKey": user_key,
            "messages": [],
            "files": []
        }

    # 3) Handle any uploaded files
    if uploaded_files and temp_container_client:
        for up_file in uploaded_files:
            if not up_file or not up_file.filename:
                continue
            file_bytes = up_file.read()
            filename = up_file.filename
            lower_name = filename.lower()

            # Extension
            if lower_name.endswith(".pdf"):
                file_ext = "pdf"
            elif lower_name.endswith((".png", ".jpg", ".jpeg")):
                file_ext = "image"
            elif lower_name.endswith(".docx"):
                file_ext = "docx"
            else:
                file_ext = "other"

            # Upload to blob
            blob_client = temp_container_client.get_blob_client(filename)
            content_settings = None
            if file_ext == "pdf":
                content_settings = ContentSettings(content_type="application/pdf")
            elif file_ext == "image":
                content_settings = ContentSettings(content_type="image/jpeg")
            elif file_ext == "docx":
                content_settings = ContentSettings(
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            blob_client.upload_blob(file_bytes, overwrite=True, content_settings=content_settings)

            base_url = temp_container_client.url
            blob_url = f"{base_url}/{filename}"

            # Possibly parse the text
            extracted_text = None
            if file_ext == "pdf":
                extracted_text = extract_text_from_pdf(file_bytes)
                messages_for_openai.append({
                    "role": "system",
                    "content": f"PDF '{filename}' uploaded.\nExtracted:\n{extracted_text[:1000]}..."
                })
            elif file_ext == "docx":
                extracted_text = extract_text_from_docx(file_bytes)
                messages_for_openai.append({
                    "role": "system",
                    "content": f"DOCX '{filename}' uploaded.\nExtracted:\n{extracted_text[:1000]}..."
                })
            elif file_ext == "image":
                vision_result = analyze_image_from_bytes(file_bytes)
                desc = "No description available."
                if vision_result.description and vision_result.description.captions:
                    desc = vision_result.description.captions[0].text
                extracted_text = f"Image AI Description: {desc}"
                messages_for_openai.append({
                    "role": "system",
                    "content": f"Image '{filename}' uploaded.\nAI says: {desc}"
                })

            chat_doc["files"].append({
                "filename": filename,
                "blobUrl": blob_url,
                "fileExt": file_ext,
                "extractedText": extracted_text or ""
            })

    # 4) Add user's message to local doc
    chat_doc["messages"].append({"role": "user", "content": user_input})

    # 5) Call AzureOpenAI
    assistant_reply = ""
    try:
        response = client.chat.completions.create(
            messages=messages_for_openai,
            model=AZURE_DEPLOYMENT_NAME
        )
        assistant_reply = response.choices[0].message.content
    except Exception as e:
        assistant_reply = f"Error calling AzureOpenAI: {e}"
        app.logger.error("OpenAI error:", exc_info=True)

    main_content = assistant_reply

    # (Optional) Filter out lines with a "download" link but no reportId
    lines = main_content.split("\n")
    filtered_lines = []
    for ln in lines:
        if re.search(r"download.*\[[^\]]+\]\([^\)]+\)", ln, flags=re.IGNORECASE) and "reportId=" not in ln:
            continue
        filtered_lines.append(ln)
    main_content = "\n".join(filtered_lines)
    main_content = re.sub(r"\n{2,}", "\n\n", main_content).strip()

    # 6) References
    references_list = []
    if "References:" in assistant_reply:
        parts = assistant_reply.split("References:")
        main_content = parts[0].strip()
        ref_section = parts[1].strip() if len(parts) > 1 else "None"
        if not ref_section.lower().startswith("none"):
            for line in ref_section.split("\n"):
                line = line.strip()
                if line.startswith("-"):
                    m = re.match(r"- \[(.*?)\]\((.*?)\): (.*)", line)
                    if m:
                        references_list.append({
                            "name": m.group(1),
                            "url": m.group(2),
                            "description": m.group(3)
                        })

    # 7) If docx link requested
    if "download://report.docx" in main_content:
        new_text = main_content.replace("download://report.docx", "").strip()
        expanded_text = generate_detailed_report(new_text)
        rid = str(uuid.uuid4())
        report_cache[rid] = expanded_text
        hyperlink = f"[Click here to download the doc](/api/generateReport?reportId={rid})"
        main_content = f"{new_text}\n\nDownloadable Report:\n{hyperlink}"

    # 8) Add assistant's reply to doc
    chat_doc["messages"].append({"role": "assistant", "content": assistant_reply})

    # 9) Write the doc back: create_item if new, otherwise replace_item
    if is_new_doc:
        container.create_item(chat_doc)
    else:
        container.upsert_item(chat_doc)

    # 10) Return JSON
    return jsonify({
        "reply": main_content,
        "references": references_list,
        "downloadUrl": None,
        "reportContent": None,
        "chatId": chat_id
    })

###############################################################################
# 8. Generate and Send Docx (GET-based)
###############################################################################
@app.route("/api/generateReport", methods=["GET"])
def generate_report_get():
    rid = request.args.get("reportId")
    if not rid:
        return "Missing reportId param", 400

    doc_text = report_cache.get(rid)
    if not doc_text:
        return "No report found for that ID or it expired.", 404

    del report_cache[rid]  # optional single-use removal

    lines = doc_text.split("\n")
    doc_title = "Generated Report"
    for idx, line in enumerate(lines):
        s = line.strip()
        if s.startswith("# "):
            doc_title = s[2:].strip()
            break
        elif s.startswith("## "):
            doc_title = s[3:].strip()
            break
        elif s.startswith("### "):
            doc_title = s[4:].strip()
            break

    doc = Document()
    doc.add_heading(doc_title, 0)

    def handle_bold(par, text):
        segments = text.split("**")
        for i, seg in enumerate(segments):
            run = par.add_run(seg)
            if i % 2 == 1:
                run.bold = True

    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif re.match(r"^-\s", stripped):
            p = doc.add_paragraph(style="List Bullet")
            handle_bold(p, stripped[2:].strip())
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            txt_part = re.sub(r"^\d+\.\s", "", stripped).strip()
            handle_bold(p, txt_part)
        else:
            par = doc.add_paragraph()
            handle_bold(par, stripped)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_title = "".join([c if c.isalnum() else "_" for c in doc_title]) or "report"
    filename = f"{safe_title}.docx"

    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

###############################################################################
# 9. Contact, Chats, Archive, Delete
###############################################################################
@app.route("/contact", methods=["POST"])
def contact_endpoint():
    data = request.get_json(force=True)
    firstName = data.get("firstName", "")
    lastName = data.get("lastName", "")
    company = data.get("company", "")
    email = data.get("email", "")
    note = data.get("note", "")

    app.logger.info(
        f"Contact form:\nName: {firstName} {lastName}\n"
        f"Company: {company}\nEmail: {email}\nNote: {note}"
    )
    if not SENDGRID_API_KEY:
        return jsonify({"status": "error", "message": "SendGrid key missing"}), 500

    try:
        sg = sendgrid.SendGridAPIClient(api_key=SENDGRID_API_KEY)
        from_email = Email("colter@mahluminnovations.com")
        to_email = To("colter@mahluminnovations.com")
        subject = f"Contact Form from {firstName} {lastName}"
        content_text = f"""
        Contact Form Submission:
        Name: {firstName} {lastName}
        Company: {company}
        Email: {email}
        Note: {note}
        """
        content = Content("text/plain", content_text)
        mail = Mail(from_email, to_email, subject, content)
        sg.client.mail.send.post(request_body=mail.get())
        return jsonify({"status": "success", "message": "Your message has been sent."}), 200
    except Exception as e:
        app.logger.error("Error sending email via SendGrid", exc_info=True)
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route("/chats", methods=["GET"])
def get_chats():
    user_key = request.args.get("userKey", "")
    if not user_key:
        return jsonify({"chats": []}), 200

    query = "SELECT * FROM c WHERE c.userKey=@userKey"
    params = [{"name": "@userKey", "value": user_key}]
    items = list(container.query_items(
        query=query,
        parameters=params,
        enable_cross_partition_query=True
    ))
    return jsonify({"chats": items}), 200

@app.route("/archiveAllChats", methods=["POST"])
def archive_all_chats():
    data = request.get_json(force=True)
    user_key = data.get("userKey", "")
    if not user_key:
        return jsonify({"error": "No userKey"}), 400

    query = "SELECT * FROM c WHERE c.userKey=@userKey"
    params = [{"name": "@userKey", "value": user_key}]
    items = list(container.query_items(query=query, parameters=params))

    for doc in items:
        doc["archived"] = True
        # STILL using upsert for archive
        container.upsert_item(doc)
    return jsonify({"success": True}), 200

@app.route("/deleteAllChats", methods=["POST"])
def delete_all_chats():
    data = request.get_json(force=True)
    user_key = data.get("userKey", "")
    if not user_key:
        return jsonify({"error": "No userKey provided"}), 400

    if AZURE_STORAGE_CONNECTION_STRING:
        from azure.storage.blob import BlobServiceClient
        try:
            bsc = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
            container_client = bsc.get_container_client("gymaitempcontainer")
        except Exception as e:
            app.logger.error("Could not connect to azure storage container:", exc_info=True)
            container_client = None
    else:
        app.logger.error("AZURE_STORAGE_CONNECTION_STRING not set; can't remove blobs.")
        container_client = None

    query = "SELECT * FROM c WHERE c.userKey=@userKey"
    params = [{"name": "@userKey", "value": user_key}]
    items = list(container.query_items(query=query, parameters=params))

    for doc in items:
        file_list = doc.get("files", [])
        if container_client and file_list:
            for f in file_list:
                try:
                    container_client.delete_blob(f["filename"])
                    app.logger.info(f"Deleted blob '{f['filename']}' from gymaitempcontainer.")
                except Exception as ex:
                    app.logger.error(f"Error deleting blob '{f['filename']}': {ex}", exc_info=True)

        container.delete_item(doc["id"], doc["userKey"])

    return jsonify({"success": True}), 200

@app.route("/renameChat", methods=["POST"])
def rename_chat():
    data = request.get_json(force=True)
    user_key = data.get("userKey", "default_user")
    chat_id = data.get("chatId", "")
    new_title = data.get("newTitle", "")

    if not chat_id or not new_title:
        return jsonify({"error": "chatId and newTitle are required"}), 400

    try:
        chat_doc = container.read_item(item=chat_id, partition_key=user_key)
        chat_doc["title"] = new_title
        container.upsert_item(chat_doc)
        return jsonify({"success": True, "message": "Title updated."}), 200
    except exceptions.CosmosResourceNotFoundError:
        return jsonify({"error": "Chat not found."}), 404
    except Exception as e:
        app.logger.error("Error renaming chat:", exc_info=True)
        return jsonify({"error": str(e)}), 500

###############################################################################
# 10. Large File Upload + Searching
###############################################################################
AZURE_SEARCH_ENDPOINT = os.environ.get("AZURE_SEARCH_ENDPOINT", "")
AZURE_SEARCH_KEY = os.environ.get("AZURE_SEARCH_KEY", "")
AZURE_SEARCH_INDEX = os.environ.get("AZURE_SEARCH_INDEX", "")

def chunk_text(text, chunk_size=1000):
    words = text.split()
    chunks = []
    current_chunk = []
    current_len = 0
    for w in words:
        current_chunk.append(w)
        current_len += len(w) + 1
        if current_len >= chunk_size:
            chunks.append(" ".join(current_chunk))
            current_chunk = []
            current_len = 0
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks

@app.route("/uploadLargeFile", methods=["POST"])
def upload_large_file():
    user_key = request.args.get("userKey", "default_user")
    if request.content_type and "multipart/form-data" in request.content_type:
        up_file = request.files.get("file")
        if not up_file:
            return jsonify({"error": "No file uploaded"}), 400
        try:
            b = up_file.read()
            fname = up_file.filename.lower()
            if fname.endswith(".pdf"):
                extracted_text = extract_text_from_pdf(b)
            elif fname.endswith(".docx"):
                extracted_text = extract_text_from_docx(b)
            else:
                return jsonify({"error": "Unsupported file type"}), 400

            chunks = chunk_text(extracted_text, 500)
            success_count = upsert_chunks_to_search(chunks, user_key)
            return jsonify({
                "status": "success",
                "message": f"Uploaded & chunked {len(chunks)} segments. {success_count} upserted."
            }), 200
        except Exception as e:
            app.logger.error("Error chunking file:", exc_info=True)
            return jsonify({"error": str(e)}), 500
    else:
        return jsonify({"error": "Please do multipart/form-data"}), 400

def upsert_chunks_to_search(chunks, user_key):
    if not (AZURE_SEARCH_ENDPOINT and AZURE_SEARCH_KEY and AZURE_SEARCH_INDEX):
        app.logger.error("Azure Search env not set.")
        return 0
    try:
        sc = SearchClient(
            endpoint=AZURE_SEARCH_ENDPOINT,
            index_name=AZURE_SEARCH_INDEX,
            credential=AzureKeyCredential(AZURE_SEARCH_KEY)
        )
        actions = []
        for ctext in chunks:
            did = f"{user_key}-{uuid.uuid4()}"
            actions.append({
                "id": did,
                "userKey": user_key,
                "content": ctext
            })
        result = sc.upload_documents(actions)
        return len(result)
    except Exception as e:
        app.logger.error("Error upserting to Search:", exc_info=True)
        return 0

@app.route("/askDoc", methods=["POST"])
def ask_doc():
    data = request.get_json(force=True) or {}
    question = data.get("question", "")
    user_key = data.get("userKey", "default_user")
    if not question:
        return jsonify({"error": "No question"}), 400

    top_chunks = search_in_azure_search(question, user_key, top_k=3)
    if not top_chunks:
        prompt_content = "No relevant documents found."
    else:
        prompt_content = "\n\n".join([f"Chunk: {c}" for c in top_chunks])

    msgs = [
        {
            "role": "system",
            "content": (
                "You are an AI that uses the following doc context. "
                "If not answered by context, say not enough info.\n"
                f"Context:\n{prompt_content}"
            )
        },
        {"role": "user", "content": question}
    ]

    try:
        resp = client.chat.completions.create(
            messages=msgs,
            model=AZURE_DEPLOYMENT_NAME
        )
        return jsonify({"answer": resp.choices[0].message.content}), 200
    except Exception as e:
        app.logger.error("Error calling OpenAI doc context:", exc_info=True)
        return jsonify({"error": str(e)}), 500

def search_in_azure_search(q, user_key, top_k=3):
    if not (AZURE_SEARCH_ENDPOINT and AZURE_SEARCH_KEY and AZURE_SEARCH_INDEX):
        app.logger.error("Azure Search not configured.")
        return []
    try:
        sc = SearchClient(
            endpoint=AZURE_SEARCH_ENDPOINT,
            index_name=AZURE_SEARCH_INDEX,
            credential=AzureKeyCredential(AZURE_SEARCH_KEY)
        )
        res = sc.search(search_text=q, filter=f"userKey eq '{user_key}'", top=top_k)
        chunks = []
        for r in res:
            c = r.get("content", "")
            if c:
                chunks.append(c)
        return chunks
    except Exception as e:
        app.logger.error("Search in Azure Search failed:", exc_info=True)
        return []

###############################################################################
# 11. Add /generateChatTitle Endpoint
###############################################################################
@app.route("/generateChatTitle", methods=["POST"])
def generate_chat_title():
    data = request.get_json(force=True) or {}
    messages = data.get("messages", [])
    model_name = data.get("model", AZURE_DEPLOYMENT_NAME)

    system_prompt = {
        "role": "system",
        "content": (
            "You are an assistant that creates short, descriptive conversation titles, "
            "3-6 words, no quotes. Return only the title as text. Avoid punctuation."
        )
    }
    messages_for_title = [system_prompt] + messages

    title_response = "Untitled Chat"
    try:
        resp = client.chat.completions.create(messages=messages_for_title, model=model_name)
        title_response = resp.choices[0].message.content.strip()
    except Exception as e:
        app.logger.error("Error calling AzureOpenAI for chat title:", exc_info=True)
        title_response = "Untitled Chat"

    return jsonify({"title": title_response})

###############################################################################
# 12. Error Handler and Main
###############################################################################
@app.errorhandler(404)
def not_found(e):
    return send_from_directory("src/public", "index.html")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080, debug=True)
